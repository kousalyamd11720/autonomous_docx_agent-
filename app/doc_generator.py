"""
doc_generator.py
-----------------
Assembles the agent's plan + step results into a polished .docx file
using python-docx.
"""
import os
import re
import uuid
from datetime import datetime
from typing import List

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from app.models import ExecutionPlan, StepResult

OUTPUT_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), "outputs")
ACCENT_COLOR = RGBColor(0x1F, 0x4E, 0x79)


def _add_title_page(doc: Document, plan: ExecutionPlan):
    title = doc.add_heading(plan.document_title, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = ACCENT_COLOR

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(plan.document_type)
    run.italic = True
    run.font.size = Pt(13)

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_p.add_run(datetime.now().strftime("%B %d, %Y"))
    date_run.font.size = Pt(10)
    date_run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    doc.add_paragraph()


def _is_table_row(line: str) -> bool:
    return line.strip().startswith("|") and line.strip().endswith("|")


def _is_separator_row(line: str) -> bool:
    # e.g. "|---|---|---|" or "| --- | :--: |"
    cells = [c.strip() for c in line.strip().strip("|").split("|")]
    return all(re.fullmatch(r":?-{2,}:?", c) for c in cells if c)


def _add_markdown_table(doc: Document, table_lines: List[str]):
    rows = [
        [c.strip() for c in ln.strip().strip("|").split("|")]
        for ln in table_lines
        if not _is_separator_row(ln)
    ]
    if not rows:
        return
    n_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.style = "Light Grid Accent 1"

    for r_idx, row in enumerate(rows):
        for c_idx in range(n_cols):
            cell_text = row[c_idx] if c_idx < len(row) else ""
            cell = table.cell(r_idx, c_idx)
            cell.text = cell_text
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10.5)
                    if r_idx == 0:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            if r_idx == 0:
                shading = OxmlElement("w:shd")
                shading.set(qn("w:fill"), "1F4E79")
                cell._tc.get_or_add_tcPr().append(shading)

    doc.add_paragraph()


def _add_assumptions(doc: Document, assumptions: List[str]):
    if not assumptions:
        return
    heading = doc.add_heading("Assumptions Made By The Agent", level=1)
    for run in heading.runs:
        run.font.color.rgb = ACCENT_COLOR
    intro = doc.add_paragraph()
    intro_run = intro.add_run(
        "Because the original request left some details unspecified, the agent "
        "made the following reasonable assumptions before generating this document:"
    )
    intro_run.italic = True
    intro_run.font.size = Pt(9.5)
    intro_run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
    for a in assumptions:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(a)
    doc.add_paragraph()


def _add_section(doc: Document, heading_text: str, content: str, used_fallback: bool):
    heading = doc.add_heading(heading_text, level=1)
    for run in heading.runs:
        run.font.color.rgb = ACCENT_COLOR

    # Split content into paragraphs/bullets/tables in a lightweight,
    # dependency-free way. Contiguous markdown-table lines ("| a | b |")
    # are rendered as a real Word table instead of plain text.
    lines = [ln.strip() for ln in content.split("\n") if ln.strip()]
    i = 0
    while i < len(lines):
        line = lines[i]
        if _is_table_row(line):
            table_lines = []
            while i < len(lines) and _is_table_row(lines[i]):
                table_lines.append(lines[i])
                i += 1
            _add_markdown_table(doc, table_lines)
            continue

        bullet_match = re.match(r"^[-*•]\s+(.*)", line)
        if bullet_match:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(bullet_match.group(1))
        else:
            p = doc.add_paragraph(line)
            p.paragraph_format.space_after = Pt(6)
        i += 1

    if used_fallback:
        note = doc.add_paragraph()
        note_run = note.add_run("(Note: generated via fallback logic — LLM was unavailable for this section.)")
        note_run.italic = True
        note_run.font.size = Pt(9)
        note_run.font.color.rgb = RGBColor(0xA0, 0x30, 0x30)

    doc.add_paragraph()


def generate_document(plan: ExecutionPlan, step_results: List[StepResult]) -> str:
    """Builds the docx file and returns its filesystem path."""
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    doc = Document()

    # Base style tweaks
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    _add_title_page(doc, plan)

    results_by_id = {r.step_id: r for r in step_results}
    for idx, step in enumerate(plan.steps):
        result = results_by_id.get(step.step_id)
        if result is None:
            continue
        _add_section(doc, result.section_heading, result.content, result.used_fallback)
        # Place assumptions right after the opening section (e.g. Executive
        # Summary) rather than before any real content -- readers see the
        # substance first, then the context needed to interpret it.
        if idx == 0:
            _add_assumptions(doc, plan.assumptions)

    filename = f"{plan.document_type.replace(' ', '_')}_{uuid.uuid4().hex[:8]}.docx"
    path = os.path.join(OUTPUT_DIR, filename)
    doc.save(path)
    return path
